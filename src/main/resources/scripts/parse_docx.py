import sys
import json
import os
import docx

def extract_docx_text(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        # 提取表格内容（可选）
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e:
        raise Exception(f"解析 DOCX 失败: {str(e)}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"success": False, "error": "缺少文件路径参数"}))
        sys.exit(1)
    file_path = sys.argv[1]
    try:
        content = extract_docx_text(file_path)
        print(json.dumps({"success": True, "content": content}, ensure_ascii=False))
    except Exception as e:
        print(json.dumps({"success": False, "error": str(e)}))